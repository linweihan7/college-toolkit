"""Build a formatted Word (.docx) minutes document from a meeting."""
from __future__ import annotations

import datetime
from pathlib import Path


def _fmt(sec: float) -> str:
    m, s = divmod(int(sec or 0), 60)
    h, m = divmod(m, 60)
    return f"{h:02d}:{m:02d}:{s:02d}" if h else f"{m:02d}:{s:02d}"


def _disp(seg: dict) -> str:
    return seg.get("edited") or seg.get("clean") or seg.get("text", "")


def build_docx(meeting: dict, out_path: Path) -> Path:
    from docx import Document

    result = meeting.get("result") or {}
    summary = result.get("summary") or {}
    names = result.get("speaker_names", {})

    doc = Document()
    doc.add_heading(meeting.get("title") or "會議記錄", 0)

    when = datetime.datetime.fromtimestamp(meeting.get("created_at", 0)).strftime("%Y-%m-%d %H:%M")
    meta = f"{when}　·　長度 {_fmt(result.get('duration', 0))}"
    if result.get("speakers"):
        labelled = "、".join(names.get(sp, sp) for sp in result["speakers"])
        meta += f"　·　與會者：{labelled}"
    doc.add_paragraph(meta)

    if summary.get("summary"):
        doc.add_heading("摘要", level=1)
        doc.add_paragraph(summary["summary"])
    if summary.get("highlights"):
        doc.add_heading("重點", level=1)
        for h in summary["highlights"]:
            doc.add_paragraph(str(h), style="List Bullet")
    if summary.get("decisions"):
        doc.add_heading("決議", level=1)
        for d in summary["decisions"]:
            doc.add_paragraph(str(d), style="List Bullet")
    if summary.get("action_items"):
        doc.add_heading("行動項目", level=1)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text = "事項", "負責人", "期限"
        for a in summary["action_items"]:
            row = table.add_row().cells
            row[0].text = a.get("task", "")
            row[1].text = a.get("owner", "")
            row[2].text = a.get("due", "")

    doc.add_heading("逐字稿", level=1)
    for s in result.get("segments", []):
        spk = names.get(s.get("speaker"), s.get("speaker")) or ""
        p = doc.add_paragraph()
        ts = p.add_run(f"[{_fmt(s['start'])}] ")
        ts.italic = True
        if spk:
            b = p.add_run(f"{spk}：")
            b.bold = True
        p.add_run(_disp(s))

    doc.save(str(out_path))
    return out_path
